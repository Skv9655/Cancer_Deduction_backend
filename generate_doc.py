from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_documentation():
    document = Document()

    # Title
    title = document.add_heading('Cancer Deduction Application Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph('\n')

    # Introduction
    heading1 = document.add_heading('1. Executive Summary', level=1)
    
    p1 = document.add_paragraph(
        "This documentation provides a deep dive into the machine learning model architecture powering the Cancer Deduction Web Application. "
        "It outlines the model selected, the rationale behind its selection, and its integration into the full-stack web ecosystem."
    )

    document.add_paragraph('\n')

    # Which model we used
    document.add_heading('2. Which Model Was Used?', level=1)
    p2 = document.add_paragraph()
    p2.add_run('Model Type: ').bold = True
    p2.add_run('XGBoost (Extreme Gradient Boosting) Classifier.\n')
    
    p2.add_run('Optimization Strategy: ').bold = True
    p2.add_run('The model was hyper-tuned using GridSearchCV to find the optimal hyperparameters, ensuring maximum accuracy and preventing overfitting. ')
    p2.add_run('The model was trained on a dataset containing 30 cellular characteristics (such as radius, texture, perimeter, area, smoothness, compactness, concavity, and symmetry).')

    document.add_paragraph('\n')

    # Why we used it
    document.add_heading('3. Why Was It Used?', level=1)
    
    document.add_paragraph(
        "XGBoost was explicitly chosen over other classical algorithms (like SVC or Random Forest) for several critical reasons:", style='List Bullet'
    )
    
    bullet1 = document.add_paragraph(style='List Bullet')
    bullet1.add_run('High Performance on Tabular Data: ').bold = True
    bullet1.add_run('XGBoost is an implementation of gradient boosted decision trees designed for speed and performance. It consistently outperforms other algorithms on structured, tabular datasets like our breast cancer feature dataset.')
    
    bullet2 = document.add_paragraph(style='List Bullet')
    bullet2.add_run('Feature Importance Extraction: ').bold = True
    bullet2.add_run('XGBoost natively calculates the importance of each feature during the tree-building process. This capability was crucial for our application, allowing us to perform feature selection and narrow down the 30 input features to the top 15 most critical ones (such as perimeter_worst, concave points_mean, and radius_worst) to streamline the user interface without heavily compromising predictive power.')
    
    bullet3 = document.add_paragraph(style='List Bullet')
    bullet3.add_run('Handling Non-Linear Relationships: ').bold = True
    bullet3.add_run('Cellular features often have complex, non-linear correlations with malignancy. XGBoost naturally captures these intricate relationships through its deep ensemble of trees.')
    
    bullet4 = document.add_paragraph(style='List Bullet')
    bullet4.add_run('Robustness to Overfitting: ').bold = True
    bullet4.add_run('Thanks to its built-in L1 and L2 regularization techniques, XGBoost is highly resistant to overfitting, making it highly reliable for medical classification tasks where false positives/negatives have severe consequences.')

    document.add_paragraph('\n')

    # Where did we use it
    document.add_heading('4. Where Was It Used?', level=1)
    
    p3 = document.add_paragraph(
        "The XGBoost model is the core analytical engine of the application. Its lifecycle and integration points are distributed across the architecture as follows:"
    )

    # 4.1 Notebook
    document.add_heading('4.1 Data Analysis & Model Training Phase (Jupyter Notebook)', level=2)
    document.add_paragraph(
        "In the initial phase (cancer_deduction.ipynb), the dataset was explored, cleaned, and processed. "
        "The XGBoost classifier was trained, hyper-tuned via GridSearchCV, and exported as a serialized pickle file (XGBoost.pkl) alongside feature importance metrics."
    )

    # 4.2 Backend
    document.add_heading('4.2 The Inference Engine (FastAPI Backend)', level=2)
    document.add_paragraph(
        "The serialized XGBoost.pkl model is loaded into the FastAPI backend (main.py) into memory upon server startup. "
        "When the backend receives a POST request at the /predict endpoint, it extracts the 15 user-provided features. "
        "Because the XGBoost model strictly requires all 30 original features to run inference, the backend dynamically reconstructs the 30-feature vector by padding the remaining 15 missing features with their pre-calculated dataset averages (loaded from feature_means.json). "
        "The backend then feeds this complete vector into the model's predict() and predict_proba() functions to return the diagnosis (Benign or Malignant) and confidence probability."
    )

    # 4.3 Frontend
    document.add_heading('4.3 The User Interface (React + Vite Frontend)', level=2)
    document.add_paragraph(
        "While the model doesn't run directly in the browser, its requirements heavily dictated the frontend UI design. "
        "The React application (App.jsx) explicitly asks the user for the Top 15 features identified by the XGBoost model's feature importance analysis. "
        "It also applies validation constraints (minimum and maximum values) based on the dataset ranges to ensure the data sent to the XGBoost backend falls within realistic biological bounds."
    )

    document.add_paragraph('\n')
    
    # Conclusion
    document.add_heading('5. Conclusion', level=1)
    document.add_paragraph(
        "By leveraging XGBoost, we built a highly accurate, robust, and performant medical prediction engine. "
        "The strategic extraction of feature importances allowed us to strike the perfect balance between clinical accuracy and user-friendly web interface design, requiring only 15 inputs from the user while the backend intelligently handles the rest."
    )

    # Save document
    document.save('Model_Documentation.docx')

if __name__ == '__main__':
    create_documentation()
